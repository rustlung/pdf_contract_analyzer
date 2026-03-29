import io
import logging

from docx import Document

from src.llm.contract_structuring_service import ContractStructuredData

logger = logging.getLogger(__name__)


class DocxGenerationError(Exception):
    """Raised when docx generation fails."""


def _value_or_default(value: str | None) -> str:
    if value is None:
        return "не указано"
    value = str(value).strip()
    return value if value else "не указано"


def _add_list_section(document: Document, title: str, values: list[str]) -> None:
    document.add_heading(title, level=2)
    if not values:
        document.add_paragraph("не указано")
        return
    for item in values:
        document.add_paragraph(item, style="List Bullet")


def generate_contract_docx(data: ContractStructuredData) -> bytes:
    logger.info("DOCX generation started")
    try:
        document = Document()
        document.add_heading("Структурированный договор (MVP)", level=1)

        document.add_heading("Общая информация", level=2)
        document.add_paragraph(f"Тип документа: {_value_or_default(data.document_type)}")
        document.add_paragraph(f"Номер договора: {_value_or_default(data.contract_number)}")
        document.add_paragraph(f"Дата договора: {_value_or_default(data.contract_date)}")

        _add_list_section(document, "Стороны", data.parties)

        document.add_heading("Предмет договора", level=2)
        document.add_paragraph(_value_or_default(data.subject))

        document.add_heading("Срок", level=2)
        document.add_paragraph(_value_or_default(data.term))

        document.add_heading("Порядок оплаты", level=2)
        document.add_paragraph(_value_or_default(data.payment_terms))

        _add_list_section(document, "Обязанности", data.obligations)
        _add_list_section(document, "Дополнительные условия", data.additional_conditions)

        document.add_heading("Примечания", level=2)
        document.add_paragraph(_value_or_default(data.notes))

        output = io.BytesIO()
        document.save(output)
        docx_bytes = output.getvalue()
    except Exception as exc:
        logger.exception("DOCX generation failed")
        raise DocxGenerationError(f"Не удалось сгенерировать DOCX: {exc}") from exc

    logger.info("DOCX generation completed successfully")
    return docx_bytes
