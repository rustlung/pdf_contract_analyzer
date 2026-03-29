import io
import logging
import re

from docx import Document
from src.api.documents.text_normalizer import normalize_extracted_text_for_docx

logger = logging.getLogger(__name__)


class DocxReconstructionError(Exception):
    """Raised when docx reconstruction fails."""


class DocxReconstructionService:
    _DOC_TITLE_RE = re.compile(r"^\s*ДОГОВОР\b.*$", re.IGNORECASE)
    _NUMBERING_RE = re.compile(r"^\s*\d+(\.\d+)*[\.)]?\s+.+$")
    _BULLET_RE = re.compile(r"^\s*[-—]\s+.+$")
    _ALL_CAPS_RE = re.compile(r"^[^a-zа-я]*[A-ZА-ЯЁ0-9][A-ZА-ЯЁ0-9\s\"'().,:;-]{5,}$")

    def generate_docx(self, raw_text: str) -> bytes:
        logger.info("DOCX reconstruction started")
        if not raw_text or not raw_text.strip():
            raise DocxReconstructionError("Пустой текст документа. Нечего реконструировать.")

        raw_lines = raw_text.splitlines()
        normalized_lines = normalize_extracted_text_for_docx(raw_text)
        logger.info(
            "DOCX reconstruction normalization applied: raw_lines=%s, normalized_paragraphs=%s",
            len(raw_lines),
            len(normalized_lines),
        )

        document = Document()
        paragraph_count = 0
        try:
            for line in normalized_lines:
                line = line.strip()
                if not line:
                    # Empty line in source text means paragraph break.
                    document.add_paragraph("")
                    paragraph_count += 1
                    continue

                if self._is_document_title(line):
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    paragraph_count += 1
                    continue

                if self._is_all_caps_heading(line):
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    paragraph_count += 1
                    continue

                if self._is_numbered_item(line):
                    document.add_paragraph(line)
                    paragraph_count += 1
                    continue

                if self._is_bullet_item(line):
                    normalized = re.sub(r"^\s*[-—]\s+", "", line)
                    document.add_paragraph(normalized, style="List Bullet")
                    paragraph_count += 1
                    continue

                document.add_paragraph(line)
                paragraph_count += 1

            buffer = io.BytesIO()
            document.save(buffer)
            content = buffer.getvalue()
        except Exception as exc:
            logger.exception("DOCX reconstruction failed")
            raise DocxReconstructionError(f"Ошибка реконструкции DOCX: {exc}") from exc

        logger.info(
            "DOCX reconstruction completed successfully: paragraphs=%s, bytes=%s",
            paragraph_count,
            len(content),
        )
        return content

    def _is_document_title(self, line: str) -> bool:
        return bool(self._DOC_TITLE_RE.match(line))

    def _is_numbered_item(self, line: str) -> bool:
        return bool(self._NUMBERING_RE.match(line))

    def _is_bullet_item(self, line: str) -> bool:
        return bool(self._BULLET_RE.match(line))

    def _is_all_caps_heading(self, line: str) -> bool:
        if len(line) > 120:
            return False
        if self._is_numbered_item(line):
            return False
        return bool(self._ALL_CAPS_RE.match(line))
